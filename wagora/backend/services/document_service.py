import io
import logging
from typing import List, Dict, Any, Optional
from pypdf import PdfReader
import docx

from services.supabase_service import db_get_brand_documents, db_get_campaign

logger = logging.getLogger("wagora-api")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using pypdf.
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            text += " | ".join(cell.text for cell in row.cells) + "\n"
    return text.strip()

def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Decodes TXT bytes to string using utf-8 or latin-1.
    """
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1").strip()

def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    Identifies the file type by extension and parses it.
    Raises ValueError if extraction yields no text or file format is unsupported.
    """
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        text = extract_text_from_docx(file_bytes)
    elif ext == "txt":
        text = extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
        
    if not text or len(text.strip()) == 0:
        raise ValueError("Document contains no readable text or is image-only")
        
    return text

async def build_ai_context(user_id: str, campaign_id: str = None) -> str:
    """
    Fetches active parsed documents for a user and compiles them into a clean string context.
    Truncates content to stay under ~2,000 tokens (8,000 characters).
    """
    docs = await db_get_brand_documents(user_id)
    
    # Sort docs by required first (business_profile, service_catalog, ideal_client, brand_voice)
    # then optional (social_proof, verbal_identity)
    priority = {
        "business_profile": 1,
        "service_catalog": 2,
        "ideal_client": 3,
        "brand_voice": 4,
        "campaign_offer": 5,
        "faq": 6,
        "objection_handling": 7,
        "social_proof": 8,
        "verbal_identity": 9
    }
    
    # Filter docs by campaign if campaign_id is provided
    relevant_docs = []
    for doc in docs:
        status = doc.get("status") or doc.get("parse_status")
        # Check active status
        if status in ["Active", "active"]:
            doc_type = doc.get("document_type", "other")
            doc_campaign_id = doc.get("campaign_id")
            
            # Account level or campaign level
            if doc_campaign_id is None or (campaign_id and str(doc_campaign_id) == str(campaign_id)):
                relevant_docs.append(doc)
                
    # Sort documents by priority
    relevant_docs.sort(key=lambda d: priority.get(d.get("document_type", "other"), 10))
    
    # Build context string
    context_parts = []
    char_count = 0
    max_chars = 8000 # 2,000 tokens proxy (8000 characters)
    
    for doc in relevant_docs:
        parsed_text = doc.get("parsed_text", "")
        if not parsed_text:
            continue
            
        doc_type_label = doc.get("document_type", "other").replace("_", " ").title()
        orig_name = doc.get("original_filename") or doc.get("name", "Document")
        
        doc_header = f"DOCUMENT: {doc_type_label} (File: {orig_name})\n---\n"
        doc_footer = "\n---\n"
        
        # Calculate available room
        remaining_chars = max_chars - char_count
        if remaining_chars <= len(doc_header) + 50:
            break # No space left for even a snippet of this document
            
        doc_body = parsed_text
        if len(doc_header) + len(doc_body) + len(doc_footer) > remaining_chars:
            # Truncate this document's text to fit
            allowed_body_len = remaining_chars - len(doc_header) - len(doc_footer) - 50
            doc_body = doc_body[:allowed_body_len] + "\n[Content truncated due to context limits]"
            
        doc_block = f"{doc_header}{doc_body}{doc_footer}"
        context_parts.append(doc_block)
        char_count += len(doc_block)
        
    # Inject Campaign Structured Fields if campaign_id is specified
    if campaign_id:
        campaign = await db_get_campaign(campaign_id)
        if campaign:
            campaign_goal = campaign.get("campaign_goal")
            target_profile = campaign.get("target_profile")
            
            structured_block = "\n--- CAMPAIGN STRUCTURAL TARGETS ---\n"
            if campaign_goal:
                structured_block += f"Goal of This Campaign: {campaign_goal}\n"
            if target_profile:
                if isinstance(target_profile, str):
                    try:
                        target_profile = json.loads(target_profile)
                    except Exception:
                        pass
                if isinstance(target_profile, dict):
                    structured_block += "Target Profile (Who This is For):\n"
                    for k, v in target_profile.items():
                        structured_block += f"  - {k.replace('_', ' ').title()}: {v}\n"
                else:
                    structured_block += f"Target Profile (Who This is For): {target_profile}\n"
            
            structured_block += "-----------------------------------\n"
            context_parts.insert(0, structured_block)
            
    if not context_parts:
        return ""
        
    return "\n".join(context_parts)
